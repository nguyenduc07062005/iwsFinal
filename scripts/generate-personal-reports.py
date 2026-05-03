from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "personal-reports"
OUT_DIR.mkdir(parents=True, exist_ok=True)

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"


COMMON_OVERVIEW = (
    "StudyVault is an AI-supported study document management system for students. "
    "The application supports account registration, email verification, login, "
    "personal document upload, file preview, folder and tag organization, search, "
    "filtering, sorting, pagination, study notes, AI summary, document Q&A, related "
    "document suggestions, and admin user management. The project follows a decoupled "
    "frontend/backend architecture: the React frontend consumes a NestJS REST API, "
    "while PostgreSQL stores both relational data and RAG-related records."
)

COMMON_ARCHITECTURE = [
    "The frontend is built with React, Vite, React Router, Tailwind CSS, Axios, Motion, and Lucide icons.",
    "The backend is built with NestJS, TypeScript, TypeORM, DTO validation, guards, and modular services.",
    "The database is PostgreSQL with pgvector support for document embeddings.",
    "Google Gemini is used for embeddings, document summaries, Q&A, mind maps, and diagrams.",
    "Authentication uses short-lived JWT access tokens, HttpOnly refresh cookies, CSRF tokens, and bcrypt password hashing.",
    "Docker Compose can run the frontend, backend, and database stack for local demonstration.",
]

REQUIREMENTS = [
    ("RESTful backend API", "NestJS controllers expose GET, POST, PATCH, and DELETE endpoints under /api."),
    ("Full CRUD", "Documents, folders, tags, notes, profile, and admin user status are implemented with appropriate HTTP methods."),
    ("Authentication and authorization", "JWT guards, refresh sessions, CSRF protection, role guards, and owner-scoped service queries protect resources."),
    ("Validation", "DTO validation, upload validation, file type checks, file size checks, and ownership checks are enforced."),
    ("Sorting and pagination", "Document lists, favorites, folder documents, admin users, and audit logs support list parameters."),
    ("Frontend API integration", "React service modules consume auth, document, folder, tag, RAG, and admin APIs."),
    ("Responsive web design", "Mobile/tablet use compact navigation; wide desktop uses expanded layouts and split document viewer."),
    ("Individual evidence", "Each personal report contains member-specific contribution, screenshots, challenges, and collaboration notes."),
]


MEMBERS = [
    {
        "idx": 1,
        "file": "Group4_Member1_Authentication_Security_Admin_Report.docx",
        "topic": "Authentication, Security, and Admin Management",
        "role": "Backend authentication, session security, authorization, and admin module owner.",
        "scope": [
            "Implemented the account lifecycle: registration, email verification, login, refresh session, logout, logout all, forgot password, reset password, change password, and profile update.",
            "Implemented secure session behavior using short-lived access tokens, HttpOnly refresh cookies, CSRF tokens, bcrypt password hashing, and session revocation.",
            "Implemented role-based admin access using UserRole.ADMIN, JwtAuthGuard, RolesGuard, and @Roles.",
            "Implemented admin APIs for stats, user listing, user lock/unlock, and audit logs.",
            "Prevented public registration from granting admin privileges and ensured normal users cannot access admin APIs.",
        ],
        "technical": [
            "AuthenticationController separates public authentication flows from protected profile and password endpoints.",
            "AuthenticationService contains the main business logic for email verification, login checks, password hashing, refresh session rotation, logout, and session revocation.",
            "User sessions store refresh token hashes and CSRF token hashes instead of raw token values.",
            "AdminController is protected by JwtAuthGuard and RolesGuard, so only admin users can call /api/admin endpoints.",
            "AdminService writes admin_audit_logs when a user is locked or unlocked, which gives the report clear database evidence.",
        ],
        "files": [
            "studyVault-backend/src/modules/authentication/authentication.controller.ts",
            "studyVault-backend/src/modules/authentication/authentication.service.ts",
            "studyVault-backend/src/modules/authentication/admin-bootstrap.service.ts",
            "studyVault-backend/src/modules/authentication/jwt/jwt.strategy.ts",
            "studyVault-backend/src/modules/authentication/roles/roles.guard.ts",
            "studyVault-backend/src/modules/admin/admin.controller.ts",
            "studyVault-backend/src/modules/admin/admin.service.ts",
            "studyVault-backend/src/database/entities/user.entity.ts",
            "studyVault-backend/src/database/entities/user-session.entity.ts",
            "studyVault-backend/src/database/entities/admin-audit-log.entity.ts",
        ],
        "api": [
            "POST /api/auth/register",
            "POST /api/auth/complete-registration",
            "POST /api/auth/login",
            "POST /api/auth/refresh",
            "POST /api/auth/logout",
            "POST /api/auth/logout-all",
            "GET /api/auth/profile",
            "PATCH /api/auth/profile",
            "PATCH /api/auth/password",
            "GET /api/admin/stats",
            "GET /api/admin/users?page=1&limit=10",
            "PATCH /api/admin/users/:id/status",
            "GET /api/admin/audit-logs?page=1&limit=10",
        ],
        "db": [
            "select id, email, role, is_active, is_email_verified, created_at from users order by created_at desc limit 10;",
            "select user_id, expires_at, revoked_at, created_at from user_sessions order by created_at desc limit 10;",
            "select admin_id, target_user_id, action, metadata, created_at from admin_audit_logs order by created_at desc limit 10;",
        ],
        "evidence": [
            "Register or complete-registration screen",
            "Login response in Postman",
            "Profile page and change-password evidence",
            "Admin dashboard stats",
            "Admin user list",
            "Lock/unlock user action",
            "Admin audit log after status change",
            "users, user_sessions, and admin_audit_logs database queries",
        ],
        "challenges": [
            ("Securing refresh sessions without exposing refresh tokens to JavaScript", "Stored refresh tokens in HttpOnly cookies and required CSRF headers for refresh/logout."),
            ("Preventing privilege escalation during registration", "Public registration always creates normal users; admin access is controlled separately."),
            ("Making admin actions auditable", "Stored admin actions in admin_audit_logs with admin id, target user id, action, metadata, and timestamp."),
            ("Invalidating risky sessions", "Revoked sessions during logout all, password changes, reset password, and admin lock flows."),
        ],
        "collaboration": [
            "Worked with Member 4 to connect Login, Profile, and Admin UI screens to the backend authentication and admin APIs.",
            "Worked with Member 2 to ensure document, folder, and tag APIs consistently use the authenticated user id.",
            "Worked with Member 3 to ensure RAG endpoints are protected and owner-scoped.",
        ],
    },
    {
        "idx": 2,
        "file": "Group4_Member2_Document_Workspace_CRUD_Report.docx",
        "topic": "Document Workspace CRUD, Folders, and Tags",
        "role": "Document workspace, folder, tag, and document organization owner.",
        "scope": [
            "Implemented upload, validation, list, detail, preview/download, rename, delete, and favorite workflows for documents.",
            "Implemented server-side pagination, search, filtering, and sorting for documents and favorites.",
            "Implemented folder CRUD, folder movement, document-folder assignment, and folder document listing.",
            "Implemented tag CRUD and document-tag assignment with owner-scoped access.",
            "Integrated workspace UI components with document, folder, and tag service modules.",
        ],
        "technical": [
            "DocumentController exposes endpoints for upload, list, detail, file preview, HTML preview, rename, delete, favorite, search, related documents, notes, and tags.",
            "DocumentService validates upload files and enforces owner checks for every document action.",
            "The database separates physical document metadata from per-user library state using document and user_documents tables.",
            "FolderController and FolderService manage each user's private folder tree and document-folder relationships.",
            "TagController and TagService manage private user tags and document-tag assignments through user_document_tags.",
        ],
        "files": [
            "studyVault-backend/src/modules/document/document.controller.ts",
            "studyVault-backend/src/modules/document/document.service.ts",
            "studyVault-backend/src/modules/folder/folder.controller.ts",
            "studyVault-backend/src/modules/folder/folder.service.ts",
            "studyVault-backend/src/modules/tag/tag.controller.ts",
            "studyVault-backend/src/modules/tag/tag.service.ts",
            "studyVault-backend/src/database/entities/document.entity.ts",
            "studyVault-backend/src/database/entities/user-document.entity.ts",
            "studyVault-backend/src/database/entities/folder.entity.ts",
            "studyVault-backend/src/database/entities/tag.entity.ts",
            "studyVault-backend/src/database/entities/user-document-tag.entity.ts",
            "studyVault-frontend/src/pages/WorkspacePage.jsx",
            "studyVault-frontend/src/components/documents/UploadModal.jsx",
            "studyVault-frontend/src/components/workspace/DocumentLibraryPanel.jsx",
            "studyVault-frontend/src/service/documentAPI.js",
            "studyVault-frontend/src/service/folderAPI.js",
            "studyVault-frontend/src/service/tagAPI.js",
        ],
        "api": [
            "POST /api/documents/upload",
            "GET /api/documents?page=1&limit=10&sortBy=createdAt&sortOrder=DESC",
            "GET /api/documents/:id",
            "PATCH /api/documents/:id",
            "DELETE /api/documents/:id",
            "POST /api/documents/:id/toggle-favorite",
            "GET /api/documents/favorites?page=1&limit=10",
            "GET /api/documents/search?q=lecture&page=1&limit=10",
            "GET /api/documents/:id/file",
            "GET /api/documents/:id/preview-html",
            "GET/POST/PATCH/DELETE /api/folders",
            "POST /api/folders/documents/add",
            "DELETE /api/folders/documents/remove",
            "GET/POST/PATCH/DELETE /api/tags",
            "PATCH /api/documents/:id/tags",
        ],
        "db": [
            "select id, title, original_name, mime_type, file_size, created_at from document order by created_at desc limit 10;",
            "select id, user_id, document_id, folder_id, display_name, is_favorite, created_at from user_documents order by created_at desc limit 10;",
            "select id, owner_id, name, parent_id, created_at from folder order by created_at desc limit 10;",
            "select id, owner_id, name, type, color, created_at from tags order by created_at desc limit 10;",
            "select user_document_id, tag_id from user_document_tags limit 20;",
        ],
        "evidence": [
            "Workspace document list",
            "Upload modal with PDF/DOCX/TXT",
            "Successful upload response",
            "Document preview or download",
            "Rename/delete/favorite document",
            "Search/filter/sort/pagination controls",
            "Folder CRUD and move document to folder",
            "Tag creation and tag assignment",
            "document, user_documents, folder, tags, and user_document_tags database queries",
        ],
        "challenges": [
            ("Upload should not fail only because AI quota is unavailable", "Stored the file and document record first, then queued RAG indexing asynchronously."),
            ("Preventing cross-user document access", "Scoped every document, folder, and tag query by ownerId or userId."),
            ("Keeping large lists scalable", "Added pagination, sorting, filtering, and search query parameters."),
            ("Supporting RESTful requirements while preserving older calls", "Kept legacy aliases and added RESTful id-based endpoints."),
        ],
        "collaboration": [
            "Worked with Member 3 because uploaded documents are the input for RAG indexing and AI assistant features.",
            "Worked with Member 4 to connect WorkspacePage, UploadModal, and DocumentLibraryPanel to API services.",
            "Worked with Member 1 to keep all workspace APIs protected by authenticated user ownership.",
        ],
    },
    {
        "idx": 3,
        "file": "Group4_Member3_AI_RAG_Document_Assistant_Report.docx",
        "topic": "AI/RAG and Document Assistant",
        "role": "AI/RAG, document assistant, summary, Q&A, and study-note owner.",
        "scope": [
            "Implemented background document indexing, chunking, embeddings, and retrieval context.",
            "Implemented Ask AI, ask history, and clear conversation for each document.",
            "Implemented English/Vietnamese summary generation, cached summary retrieval, and regenerate behavior.",
            "Implemented related documents, mind map, diagram endpoints, and study-note workflows.",
            "Configured Gemini model and fallback behavior for more reliable demo under free-tier limits.",
        ],
        "technical": [
            "RagController exposes summary, cached summary, ask, ask history, clear history, mind map, and diagram APIs.",
            "RagIndexingService ensures uploaded documents can be chunked and indexed for retrieval.",
            "RagQuestionAnsweringService builds document-scoped context and stores Q&A records in document_ask_history.",
            "RagSummaryService separates GET cached summary from POST generate summary to avoid unnecessary AI calls.",
            "GeminiService uses configured model ids and fallback models to reduce failures when a model is unavailable.",
            "DocumentViewer integrates Study, Summary, Ask AI, and Related tabs into one assistant panel.",
        ],
        "files": [
            "studyVault-backend/src/modules/rag/rag.controller.ts",
            "studyVault-backend/src/modules/rag/rag.service.ts",
            "studyVault-backend/src/modules/rag/services/rag-indexing.service.ts",
            "studyVault-backend/src/modules/rag/services/rag-question-answering.service.ts",
            "studyVault-backend/src/modules/rag/services/rag-summary.service.ts",
            "studyVault-backend/src/modules/rag/services/rag-search.service.ts",
            "studyVault-backend/src/modules/rag/services/rag-mind-map.service.ts",
            "studyVault-backend/src/modules/rag/services/rag-artifact-cache.service.ts",
            "studyVault-backend/src/common/llm/gemini.service.ts",
            "studyVault-backend/src/database/entities/chunks.entity.ts",
            "studyVault-backend/src/database/entities/document-ask-history.entity.ts",
            "studyVault-backend/src/database/entities/study-note.entity.ts",
            "studyVault-frontend/src/pages/DocumentViewer.jsx",
            "studyVault-frontend/src/service/ragAPI.js",
        ],
        "api": [
            "POST /api/rag/documents/:documentId/summary",
            "GET /api/rag/documents/:documentId/summary?language=en",
            "GET /api/rag/documents/:documentId/summary?language=vi",
            "POST /api/rag/documents/:documentId/ask",
            "GET /api/rag/documents/:documentId/ask/history",
            "DELETE /api/rag/documents/:documentId/ask/history",
            "POST /api/rag/documents/:documentId/mindmap",
            "POST /api/rag/documents/:documentId/diagram",
            "GET /api/documents/:id/related?limit=6",
            "GET/POST/PATCH/DELETE document note endpoints",
        ],
        "db": [
            "select id, document_id, chunk_index, token_count, embedding_model, created_at from chunks order by created_at desc limit 10;",
            "select id, user_id, document_id, question, answer, created_at from document_ask_history order by created_at desc limit 10;",
            "select id, user_id, user_document_id, content, created_at, updated_at from study_notes order by updated_at desc limit 10;",
        ],
        "evidence": [
            "Document Assistant panel",
            "Summary tab before and after generation",
            "English and Vietnamese summaries",
            "Ask AI conversation",
            "Ask history after reload",
            "Clear conversation result",
            "Related documents tab",
            "Study notes create/update/delete",
            "chunks, document_ask_history, and study_notes database queries",
        ],
        "challenges": [
            ("Gemini free-tier quota and model availability", "Checked available model ids and configured fallback models."),
            ("Avoiding accidental AI generation when loading a cached summary", "Separated cached GET summary from generate POST summary."),
            ("Preventing AI from reading another user's documents", "Loaded RAG context only after owner checks passed."),
            ("Making the assistant usable on iPad", "Coordinated responsive DocumentViewer changes so preview and assistant stack below wide desktop."),
        ],
        "collaboration": [
            "Worked with Member 2 because document upload and extraction feed the indexing pipeline.",
            "Worked with Member 4 to make Document Assistant responsive and connect ragAPI endpoints.",
            "Worked with Member 1 to keep RAG APIs protected by JWT and user ownership.",
        ],
    },
    {
        "idx": 4,
        "file": "Group4_Member4_Frontend_Responsive_Docs_Testing_Report.docx",
        "topic": "Frontend UX, Responsive Design, Integration, Documentation, and Testing",
        "role": "Frontend UX, responsive layout, API integration, documentation, and verification owner.",
        "scope": [
            "Implemented and maintained public, auth, protected app, and document detail route groups.",
            "Implemented AppShell, AuthLayout, DetailLayout, ShellHeader, and responsive navigation behavior.",
            "Integrated pages with auth, document, folder, tag, RAG, and admin API service modules.",
            "Fixed iPad header overlap and tablet/mobile DocumentViewer overflow issues.",
            "Prepared documentation, report templates, evidence plan, responsive QA checklist, and frontend verification evidence.",
        ],
        "technical": [
            "router.jsx separates landing, guest-only authentication pages, protected workspace pages, admin page, profile page, and document detail pages.",
            "ShellHeader keeps a compact top bar and bottom navigation below 1024px, then switches to centered desktop navigation at 1024px and wider.",
            "DetailLayout and DocumentViewer allow scrolling and stacked preview/assistant below 1280px, while wide desktop uses split preview/assistant layout.",
            "AuthLayout uses min-h-dvh so auth forms remain usable on shorter screens and when a software keyboard is open.",
            "Frontend service modules keep API calls organized by feature instead of scattering HTTP calls inside pages.",
            "README and docs were updated with route map, component map, responsive QA targets, and personal report guidance.",
        ],
        "files": [
            "studyVault-frontend/src/app/router.jsx",
            "studyVault-frontend/src/layouts/AppShell.jsx",
            "studyVault-frontend/src/layouts/AuthLayout.jsx",
            "studyVault-frontend/src/layouts/DetailLayout.jsx",
            "studyVault-frontend/src/components/navigation/ShellHeader.jsx",
            "studyVault-frontend/src/pages/Landing.jsx",
            "studyVault-frontend/src/pages/WorkspacePage.jsx",
            "studyVault-frontend/src/pages/Favorites.jsx",
            "studyVault-frontend/src/pages/Profile.jsx",
            "studyVault-frontend/src/pages/Admin.jsx",
            "studyVault-frontend/src/pages/DocumentViewer.jsx",
            "studyVault-frontend/src/service/*.js",
            "README.md",
            "studyVault-frontend/README.md",
            "docs/frontend-route-map.md",
            "docs/frontend-component-map.md",
            "docs/member-report-evidence-plan.md",
        ],
        "api": [
            "Browser Network: GET /api/auth/profile",
            "Browser Network: GET /api/documents",
            "Browser Network: POST /api/documents/upload",
            "Browser Network: GET /api/documents/:id/file",
            "Browser Network: POST /api/rag/documents/:id/ask",
            "Browser Network: GET /api/admin/stats",
            "Browser Network: GET /api/admin/users",
        ],
        "db": [
            "Use browser Network screenshots and terminal verification as primary evidence.",
            "Optional shared DB evidence: users, document, user_documents, tags, and admin_audit_logs to prove the UI creates real data.",
        ],
        "evidence": [
            "Landing page desktop",
            "Login/Register mobile",
            "Workspace at iPad portrait 768x1024",
            "Workspace at iPad landscape 1024x768",
            "Document Viewer stacked layout on tablet/mobile",
            "Document Viewer split layout on desktop",
            "Favorites/Profile/Admin responsive screens",
            "Browser Network API calls",
            "npm run lint, npm test, and npm run build terminal outputs",
        ],
        "challenges": [
            ("iPad header overlap caused by desktop navigation starting too early", "Moved centered desktop navigation to the lg breakpoint and kept bottom navigation through tablet widths."),
            ("Document Viewer content was clipped on smaller screens", "Allowed page scrolling below xl and used stacked preview/assistant layout until wide desktop."),
            ("Personal reports could become too similar", "Created four separate reports with different contribution, evidence, challenge, and collaboration sections."),
            ("Responsive evidence was required by the project brief", "Added viewport checklist for 375x667, 768x1024, 1024x768, 1366x768, and 1440x900."),
        ],
        "collaboration": [
            "Worked with Member 1 to connect Login, Profile, and Admin UI screens with protected APIs.",
            "Worked with Member 2 to integrate workspace upload, document list, folder, and tag controls.",
            "Worked with Member 3 to make Document Assistant responsive and connect summary/ask/history endpoints.",
        ],
    },
]


def set_run(run, size=None, bold=None, italic=None):
    run.font.name = FONT
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.color.rgb = BLACK
        style.font.size = Pt(11)


def add_paragraph(doc, text="", bold=False, italic=False, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run(run, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    size = 14 if level == 1 else 12.5
    return add_paragraph(doc, text, bold=True, align=None).runs[0].font.__setattr__("size", Pt(size))


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run(run)


def add_numbers(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run(run)


def set_black_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_black_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(str(header))
        set_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            run = cell.paragraphs[0].add_run(str(value))
            set_run(run)
    doc.add_paragraph("")
    return table


def add_placeholder(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, bold=True, italic=True)


def enforce_black(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run(run)
    for table in doc.tables:
        set_black_borders(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run(run)


def cover(doc, member):
    title = "PERSONAL PROJECT REPORT" if member else "REPORT TEMPLATE"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_run(run, size=22, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("StudyVault - AI Study Document Management System")
    set_run(run, size=16, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("IWS Spring 2026 - Group 4")
    set_run(run, size=12, bold=True)
    doc.add_paragraph("")
    if member:
        rows = [
            ("Group", "Group 4"),
            ("Student name", "[FULL NAME]"),
            ("Student ID", "[STUDENT ID]"),
            ("Member number", f"Member {member['idx']}"),
            ("Individual topic", member["topic"]),
            ("Role", member["role"]),
            ("Submission date", "[DATE]"),
        ]
    else:
        rows = [
            ("Group", "Group 4"),
            ("Purpose", "Reference only"),
            ("Instruction", "Use the four individual reports in docs/personal-reports for submission."),
            ("Format", "Plain black text and black table borders. No blue heading text or blue table shading."),
        ]
    add_table(doc, ["Field", "Content"], rows)
    add_paragraph(
        doc,
        "Shared project overview sections may match across members, but contribution, evidence, challenges, and collaboration must be individual.",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def common_sections(doc):
    add_heading(doc, "1. Project Overview")
    add_paragraph(doc, COMMON_OVERVIEW)
    add_paragraph(
        doc,
        "The project targets the IWS final project requirements: RESTful backend API, client-side API consumption, CRUD, authentication and authorization, validation, sorting, pagination, and responsive web design.",
    )

    add_heading(doc, "2. Main Functions")
    add_table(
        doc,
        ["Function group", "Detailed description"],
        [
            ("Authentication", "Registration, email verification, login, refresh session, logout, logout all, password reset, password change, and profile update."),
            ("Document workspace", "Upload PDF/DOCX/TXT, validate files, preview/download documents, rename, delete, favorite, search, filter, sort, and paginate."),
            ("Folders and tags", "Folder CRUD, folder movement, document-folder assignment, tag CRUD, and document-tag assignment."),
            ("AI/RAG", "Document indexing, chunks, embeddings, summary EN/VI, cached summary, Ask AI, ask history, related documents, mind map, and diagram."),
            ("Admin", "Stats, user list, lock/unlock normal users, audit logs, and admin-only authorization."),
            ("Responsive UI", "Mobile/tablet compact navigation, iPad-safe layout, and wide desktop split document viewer."),
        ],
    )

    add_heading(doc, "3. Technology Stack")
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Frontend", "React 19, Vite, Tailwind CSS, React Router, Axios, Motion, Lucide", "UI, routing, API calls, responsive layout, and user interaction."),
            ("Backend", "NestJS, TypeScript, TypeORM, class-validator", "REST API, validation, guards, services, and business logic."),
            ("Database", "PostgreSQL, pgvector", "Relational data and vector embedding storage."),
            ("Authentication", "JWT, HttpOnly cookie, CSRF token, bcrypt", "Secure account and session handling."),
            ("AI", "Google Gemini", "Embeddings, summary, Q&A, mind map, and diagram generation."),
            ("Runtime", "Docker Compose, npm scripts", "Local demo, build, test, and readiness verification."),
        ],
    )

    add_heading(doc, "4. Architecture and Database Overview")
    add_bullets(doc, COMMON_ARCHITECTURE)
    add_table(
        doc,
        ["Database table", "Purpose"],
        [
            ("users", "Account data, role, active flag, email verification, and password reset fields."),
            ("user_sessions", "Refresh session hash, CSRF hash, expiry, and revoked state."),
            ("document", "Uploaded file metadata and storage reference."),
            ("user_documents", "Per-user document library entry, display name, folder, and favorite state."),
            ("folder", "Private folder tree per user."),
            ("tags / user_document_tags", "Private labels and document-tag assignments."),
            ("chunks", "Text chunks and embedding metadata for RAG."),
            ("document_ask_history", "AI question-answer history per user and document."),
            ("study_notes", "Study notes attached to user documents."),
            ("admin_audit_logs", "Audit evidence for admin user status changes."),
        ],
    )

    add_heading(doc, "5. Requirement Mapping")
    add_table(doc, ["IWS requirement", "StudyVault implementation"], REQUIREMENTS)

    add_heading(doc, "6. Shared Verification Plan")
    add_paragraph(doc, "These verification items may be shared across reports, but each member must still include evidence specific to their own work.")
    add_bullets(
        doc,
        [
            "Frontend: npm run lint, npm test, npm run build.",
            "Backend: npm run build, npm test, npm run test:e2e.",
            "API: capture request and response screenshots from Postman or Swagger.",
            "Database: capture queries related to the member's module.",
            "Responsive: capture 375x667, 768x1024, 1024x768, 1366x768, and 1440x900.",
        ],
    )


def individual_sections(doc, member):
    add_heading(doc, "7. Individual Contribution")
    add_paragraph(doc, member["role"], bold=True)
    add_paragraph(
        doc,
        "This is the most important part of the personal report. It must be rewritten by the assigned member and supported with screenshots.",
    )
    add_bullets(doc, member["scope"])

    add_heading(doc, "8. Assigned Tasks")
    add_table(doc, ["Assigned task", "Required proof"], [(task, "Insert matching UI/API/DB/code/test evidence.") for task in member["scope"]])

    add_heading(doc, "9. Detailed Technical Contribution")
    for item in member["technical"]:
        add_paragraph(doc, item)
    add_heading(doc, "9.1 Related Files and Modules", 2)
    add_bullets(doc, member["files"])

    add_heading(doc, "10. Evidence Screenshots")
    add_paragraph(doc, "Each figure must include a caption explaining what feature, endpoint, database query, or test command it proves.")
    add_heading(doc, "10.1 UI Evidence", 2)
    add_bullets(doc, member["evidence"])
    add_heading(doc, "10.2 API Evidence", 2)
    add_bullets(doc, member["api"])
    add_heading(doc, "10.3 Database or Test Evidence", 2)
    add_bullets(doc, member["db"])
    for index, evidence in enumerate(member["evidence"], 1):
        add_placeholder(doc, f"[INSERT FIGURE {index}: {evidence}]")
        add_paragraph(doc, f"Figure {index}. {evidence}. Caption: [write the route/API/table/command and result shown in the screenshot].")

    add_heading(doc, "11. Challenges and Solutions")
    add_table(doc, ["Challenge", "Solution"], member["challenges"])
    add_paragraph(doc, "[Add one personal paragraph explaining the hardest technical issue and how it was verified.]")

    add_heading(doc, "12. Collaboration With Other Members")
    add_bullets(doc, member["collaboration"])
    add_paragraph(doc, "[Replace Member 1/2/3/4 with real names and describe exactly what was done together.]")

    add_heading(doc, "13. Personal Reflection")
    add_paragraph(doc, "[Write 1-2 paragraphs about what you learned, which part you are most confident in, remaining limitations, and what you would improve with more time.]")

    add_heading(doc, "14. Personal Submission Checklist")
    add_numbers(
        doc,
        [
            "Replace [FULL NAME], [STUDENT ID], and [DATE].",
            "Insert UI/API/DB/test screenshots for your own assigned work.",
            "Write the individual contribution in your own words.",
            "Describe your own challenges and solutions.",
            "State who you collaborated with and what you did together.",
            "Remove or hide API keys, passwords, full access tokens, cookies, and real .env content.",
        ],
    )


def make_member_report(member):
    doc = Document()
    configure_document(doc)
    cover(doc, member)
    common_sections(doc)
    individual_sections(doc, member)
    enforce_black(doc)
    path = OUT_DIR / member["file"]
    path = path.with_name(f"{path.stem}_EN_Black{path.suffix}")
    doc.save(path)
    return path


def make_reference_template():
    doc = Document()
    configure_document(doc)
    cover(doc, None)
    common_sections(doc)
    add_heading(doc, "7. Individual Reports Required")
    add_paragraph(
        doc,
        "This file is a reference template only. For submission, use the four separate personal reports in docs/personal-reports.",
    )
    add_table(doc, ["Member", "Report file"], [(f"Member {m['idx']}", m["file"]) for m in MEMBERS])
    enforce_black(doc)
    path = ROOT / "docs" / "StudyVault_Final_Report_Template_EN_Black.docx"
    doc.save(path)
    return path


def main():
    paths = [make_member_report(member) for member in MEMBERS]
    paths.append(make_reference_template())
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
